import os
import logging
from pathlib import Path
from typing import Optional, Tuple, Dict, Any
import chardet
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import mammoth
from PIL import Image
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdfminer.layout import LAParams
import cv2
import numpy as np
from transformers import pipeline
from config.settings import settings

logger = logging.getLogger(__name__)

class DocumentReader:
    """Advanced document reader with improved extraction and OCR capabilities"""
    
    def __init__(self, enable_ocr: bool = None, ocr_threshold: float = None):
        self.enable_ocr = enable_ocr if enable_ocr is not None else settings.ENABLE_OCR
        self.ocr_threshold = ocr_threshold if ocr_threshold is not None else settings.OCR_CONFIDENCE_THRESHOLD
        self.ocr_lang = settings.OCR_LANGUAGE
        self.ocr_config = settings.OCR_CONFIG
        self._init_ocr()
        
    def _init_ocr(self):
        """Initialize OCR settings and models"""
        try:
            # Configure Tesseract
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
            
            # Initialize OCR pipeline
            self.ocr_pipeline = pipeline("document-question-answering", 
                                      model="microsoft/layoutlmv3-base",
                                      device=0 if settings.USE_GPU else -1)
            
            # Set OCR parameters
            self.ocr_params = {
                'lang': self.ocr_lang,
                'config': self.ocr_config,
                'nice': 0,
                'timeout': 30
            }
        except Exception as e:
            logger.error(f"Error initializing OCR: {e}")
            self.ocr_pipeline = None
    
    def _get_file_type(self, file_path: str) -> str:
        """Detect file type using file signatures and extension"""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)  # Read first 8 bytes for signature
                
                # Check file signatures
                if header.startswith(b'%PDF-'):
                    return 'application/pdf'
                elif header.startswith(b'PK\x03\x04'):
                    return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif header.startswith(b'\xD0\xCF\x11\xE0'):
                    return 'application/msword'  # DOC file signature
                
                # Fallback to extension-based detection
                ext = Path(file_path).suffix.lower()
                mime_types = {
                    '.pdf': 'application/pdf',
                    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    '.doc': 'application/msword',
                    '.txt': 'text/plain',
                    '.rtf': 'application/rtf'
                }
                return mime_types.get(ext, 'application/octet-stream')
                
        except Exception as e:
            logger.error(f"Error detecting file type for {file_path}: {e}")
            return 'application/octet-stream'
    
    def detect_encoding(self, file_path: str) -> str:
        """Detect file encoding with improved accuracy"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)
                result = chardet.detect(raw_data)
                confidence = result['confidence']
                
                # If confidence is low, try reading more data
                if confidence < 0.8:
                    f.seek(0)
                    raw_data = f.read(50000)
                    result = chardet.detect(raw_data)
                
                return result['encoding'] or 'utf-8'
        except Exception as e:
            logger.error(f"Error detecting encoding for {file_path}: {e}")
            return 'utf-8'
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR results"""
        try:
            # Convert to numpy array
            img_array = np.array(image)
            
            # Convert to grayscale
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # Apply adaptive thresholding
            thresh = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # Denoise
            denoised = cv2.fastNlMeansDenoising(thresh)
            
            # Convert back to PIL Image
            return Image.fromarray(denoised)
        except Exception as e:
            logger.error(f"Error preprocessing image: {e}")
            return image
    
    def read_pdf_with_ocr(self, file_path: str) -> Tuple[str, bool]:
        """Read PDF with improved OCR and text extraction"""
        text = ""
        used_ocr = False
        
        try:
            # Try multiple PDF extraction methods
            extraction_methods = [
                self._extract_with_pdfplumber,
                self._extract_with_pdfminer,
                self._extract_with_pypdf2
            ]
            
            for method in extraction_methods:
                try:
                    text = method(file_path)
                    if len(text.strip()) > 100:
                        logger.info(f"Successfully extracted text using {method.__name__}")
                        return text, False
                except Exception as e:
                    logger.warning(f"Method {method.__name__} failed: {e}")
                    continue
            
            # If all methods fail or produce poor results, use OCR
            if self.enable_ocr and len(text.strip()) < 100:
                logger.info(f"Using OCR for {file_path}")
                images = convert_from_path(file_path)
                for image in images:
                    # Preprocess image
                    processed_image = self._preprocess_image(image)
                    
                    # Try advanced OCR first
                    if self.ocr_pipeline:
                        try:
                            ocr_result = self.ocr_pipeline(processed_image)
                            if ocr_result and 'text' in ocr_result:
                                text += ocr_result['text'] + "\n"
                                continue
                        except Exception as e:
                            logger.warning(f"Advanced OCR failed: {e}")
                    
                    # Fallback to Tesseract
                    ocr_text = pytesseract.image_to_string(
                        processed_image,
                        **self.ocr_params
                    )
                    text += ocr_text + "\n"
                used_ocr = True
                
            return text, used_ocr
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return "", False
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber"""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    
    def _extract_with_pdfminer(self, file_path: str) -> str:
        """Extract text using pdfminer"""
        laparams = LAParams(
            line_margin=0.5,
            word_margin=0.1,
            char_margin=2.0,
            boxes_flow=0.5,
            detect_vertical=True
        )
        return pdfminer_extract_text(file_path, laparams=laparams)
    
    def _extract_with_pypdf2(self, file_path: str) -> str:
        """Extract text using PyPDF2"""
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def read_docx(self, file_path: str) -> str:
        """Read DOCX file with improved extraction"""
        try:
            # Try mammoth first for better formatting preservation
            try:
                with open(file_path, 'rb') as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    if result.value:
                        return result.value
            except Exception as e:
                logger.warning(f"Mammoth extraction failed: {e}")
            
            # Fallback to python-docx
            doc = Document(file_path)
            
            # Extract text with formatting information
            text_parts = []
            for paragraph in doc.paragraphs:
                # Get paragraph text with basic formatting
                text = paragraph.text
                if text.strip():
                    # Add formatting markers
                    if paragraph.style.name.startswith('Heading'):
                        text = f"\n{text}\n"
                    text_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def read_document(self, file_path: str, max_chars: int = 50000) -> Tuple[str, bool]:
        """Read document with improved format detection and handling"""
        try:
            # Detect file type
            file_type = self._get_file_type(file_path)
            
            # Read based on file type
            if file_type == 'application/pdf':
                text, used_ocr = self.read_pdf_with_ocr(file_path)
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                text = self.read_docx(file_path)
                used_ocr = False
            elif file_type == 'application/msword':
                # Try converting DOC to DOCX first
                try:
                    import win32com.client
                    import os
                    import tempfile
                    import pythoncom
                    import time
                    
                    # Initialize COM for this thread
                    pythoncom.CoInitialize()
                    
                    # Create a temporary file for the DOCX
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                        temp_docx = temp_file.name
                    
                    try:
                        # Convert DOC to DOCX using Word
                        word = None
                        try:
                            word = win32com.client.Dispatch('Word.Application')
                            word.Visible = False
                            word.DisplayAlerts = False
                            
                            doc = word.Documents.Open(os.path.abspath(file_path))
                            doc.SaveAs(os.path.abspath(temp_docx), 16)  # 16 represents DOCX format
                            doc.Close()
                            
                            # Give Word time to complete operations
                            time.sleep(1)
                            
                        finally:
                            if word:
                                try:
                                    word.Quit()
                                except:
                                    pass
                                # Give Word time to quit
                                time.sleep(1)
                        
                        # Read the converted DOCX
                        text = self.read_docx(temp_docx)
                        used_ocr = False
                        
                    finally:
                        # Clean up temporary file
                        try:
                            os.unlink(temp_docx)
                        except:
                            pass
                        
                except Exception as e:
                    logger.error(f"Error converting DOC to DOCX: {e}")
                    # Fall back to antiword if conversion fails
                    try:
                        text = self.read_doc(file_path)
                        used_ocr = False
                    except Exception as e:
                        logger.error(f"Error reading DOC with antiword: {e}")
                        text = ""
                        used_ocr = False
            else:
                logger.error(f"Unsupported file type: {file_type}")
                return "", False
            
            # Truncate if needed
            if len(text) > max_chars:
                text = text[:max_chars]
            
            return text, used_ocr
            
        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            return "", False
    
    def read_doc(self, file_path: str) -> str:
        """Read DOC file using antiword with full path"""
        try:
            import subprocess
            antiword_path = r'C:\antiword\antiword.exe'
            result = subprocess.run(
                [antiword_path, file_path], 
                capture_output=True, 
                text=True,
                encoding='utf-8',
                errors='replace'  # Handle encoding errors gracefully
            )
            if result.returncode == 0:
                return result.stdout
            else:
                logger.error(f"antiword failed for {file_path}: {result.stderr}")
                return ""
        except Exception as e:
            logger.error(f"Error reading DOC {file_path} with antiword: {e}")
            return "" 